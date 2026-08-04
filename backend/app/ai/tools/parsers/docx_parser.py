from io import BytesIO
import docx
from app.ai.tools.parsers.base import BaseDocumentParser


class DocxParser(BaseDocumentParser):
    """
    Parses Microsoft Word (.docx) files.
    """

    async def parse(self, file_bytes: bytes) -> str:
        stream = BytesIO(file_bytes)
        doc = docx.Document(stream)

        # Extract text from paragraphs
        paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_cells = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_cells:
                    # Tab-separate cells within rows
                    table_text.append("\t".join(row_cells))

        all_text = paragraphs_text + table_text
        return "\n".join(all_text)
