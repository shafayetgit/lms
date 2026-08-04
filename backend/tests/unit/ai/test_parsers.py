import io
from unittest.mock import AsyncMock, patch
import docx
import fitz
import pytest
from app.ai.tools.parsers.docx_parser import DocxParser
from app.ai.tools.parsers.router import ParserRouter


@pytest.mark.asyncio
async def test_docx_parser():
    doc = docx.Document()
    doc.add_paragraph("Hello from docx paragraph.")

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header1"
    table.cell(0, 1).text = "Header2"
    table.cell(1, 0).text = "Value1"
    table.cell(1, 1).text = "Value2"

    stream = io.BytesIO()
    doc.save(stream)
    file_bytes = stream.getvalue()

    parser = DocxParser()
    extracted = await parser.parse(file_bytes)

    assert "Hello from docx paragraph." in extracted
    assert "Header1\tHeader2" in extracted
    assert "Value1\tValue2" in extracted


@pytest.mark.asyncio
async def test_parser_router_pdf_ocr_fallback():
    # Create empty PDF bytes
    doc = fitz.open()
    doc.new_page()
    stream = io.BytesIO()
    doc.save(stream)
    pdf_bytes = stream.getvalue()
    doc.close()

    router = ParserRouter()

    # Mock the ImageVisionParser parse method to return OCRed text
    with patch.object(
        router.parsers[".jpg"], "parse", new_callable=AsyncMock
    ) as mock_vision_parse:
        mock_vision_parse.return_value = "Extracted OCR text."

        extracted = await router.parse_file("scanned.pdf", pdf_bytes)

        assert extracted == "Extracted OCR text."
        mock_vision_parse.assert_called_once()


@pytest.mark.asyncio
async def test_parser_router_pdf_ocr_fallback_capping():
    # Create a 12-page empty PDF
    doc = fitz.open()
    for _ in range(12):
        doc.new_page()
    stream = io.BytesIO()
    doc.save(stream)
    pdf_bytes = stream.getvalue()
    doc.close()

    router = ParserRouter()

    with patch.object(
        router.parsers[".jpg"], "parse", new_callable=AsyncMock
    ) as mock_vision_parse:
        mock_vision_parse.return_value = "Page Text"

        await router.parse_file("scanned.pdf", pdf_bytes)

        # Verify it was capped to 10 pages
        assert mock_vision_parse.call_count == 10
